import sys
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.enum.text import WD_BREAK, WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE

def create_field_code(run, field_text):
    """构建 Word 底层域代码"""
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = f" {field_text} "
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

def setup_toc_styles(doc):
    """
    黑客魔法：提前劫持并篡改 Word 的内置 TOC 样式
    满足：固定值20磅，两端对齐，中英文字体Times New Roman，精确缩进
    """
    # 缩进计算：小四号字是 12 磅。1个字符缩进就是 12 磅，2个就是 24 磅。
    styles_config = {
        'TOC 1': {'size': 14, 'indent': 0},   # 一级目录：14磅(四号)，无缩进
        'TOC 2': {'size': 12, 'indent': 12},  # 二级目录：12磅(小四)，缩进 12 磅
        'TOC 3': {'size': 12, 'indent': 24}   # 三级目录：12磅(小四)，缩进 24 磅
    }

    for style_name, config in styles_config.items():
        try:
            style = doc.styles[style_name]
        except KeyError:
            # 如果文档太干净，没有这个样式，我们就强行塞一个进去
            style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)

        # 1. 字体与字号设定
        style.font.name = 'Times New Roman'
        # 必须加上这句底层 XML 强制替换，否则 Word 可能会擅自把汉字变成宋体
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        style.font.size = Pt(config['size'])

        # 2. 段落行距：固定值 20 磅
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        style.paragraph_format.line_spacing = Pt(20)

        # 3. 对齐方式：两端对齐
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # 4. 精确左缩进
        if config['indent'] > 0:
            style.paragraph_format.left_indent = Pt(config['indent'])
            
        # 注意：Word 的 TOC 样式自带“页码右对齐”的制表符，我们保留它的默认属性即可

def prepend_directories(doc_path, output_path):
    print(f"正在注入带有顶级论文排版规范的目录...\n- 目标文件: {doc_path}")
    doc = Document(doc_path)
    
    # 第一步：提前篡改文档的 TOC 样式字典
    setup_toc_styles(doc)
    
    if not doc.paragraphs:
        print("错误：文档是空的！")
        sys.exit(1)
    first_p = doc.paragraphs[0]
    
    def insert_directory(title_text, field_code, need_page_break=True):
        # 1. 插入大标题 (严格要求：居中，单倍行距)
        title_p = first_p.insert_paragraph_before("")
        title_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        
        # 标题文字格式 (严格要求：Times New Roman，16磅/三号，加粗)
        run_title = title_p.add_run(title_text)
        run_title.bold = True
        run_title.font.name = 'Times New Roman'
        run_title._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        run_title.font.size = Pt(16)
        
        # 2. 插入域代码 + 分页符
        field_p = first_p.insert_paragraph_before("")
        run_field = field_p.add_run()
        create_field_code(run_field, field_code)
        if need_page_break:
            field_p.add_run().add_break(WD_BREAK.PAGE)

    # 按照正序在页面最前方倒序插入，保证最终顺序
    insert_directory("Table of Contents", 'TOC \\o "1-3" \\h \\z \\u')
    insert_directory("List of Figures", 'TOC \\h \\z \\c "Figure"')
    insert_directory("List of Tables", 'TOC \\h \\z \\c "Table"')

    doc.save(output_path)
    print(f"排版级注入完成！请打开 {output_path} 执行全选更新。")

if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("用法: python prepend_toc.py <输入文档.docx> <输出文档.docx>")
        sys.exit(1)
        
    prepend_directories(sys.argv[1], sys.argv[2])
